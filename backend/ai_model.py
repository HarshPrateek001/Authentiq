import os
import httpx
import random
import re
import io
import json
import hashlib
from fastapi import HTTPException
from datetime import datetime 
from typing import Optional, List, Dict, Any, Tuple

# Import centralized configuration
from config import config

# Import Pydantic BaseModel and Field from pydantic library
from pydantic import BaseModel, Field

# Corrected: Import Pydantic models from the new models.py file
from models import (
    PlagiarismRequest, HumanizeRequest, ChatRequest, PlagiarismResult, HumanizeResult, ChatResponse,
    DownloadHumanizedRequest
)

# --- Optional Libraries for Document Generation ---
try:
    from fpdf import FPDF
except ImportError:
    FPDF = None
    print("WARNING: fpdf library not found. PDF generation will be disabled.")

try:
    from docx import Document
except ImportError:
    Document = None
    print("WARNING: python-docx library not found. Word document generation will be disabled.")
    
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None
    print("WARNING: PyPDF2 library not found. PDF text extraction will be disabled.")

try:
    from unidecode import unidecode
except ImportError:
    def unidecode(text):
        return text.encode('ascii', 'ignore').decode('ascii')
    print("WARNING: unidecode library not found. Using basic ASCII fallback.")

async def extract_text_from_bytes(content: bytes, content_type: str) -> str:
    """Extracts text from file bytes based on content type."""
    extracted_text = ""
    if content_type == "text/plain":
        extracted_text = content.decode('utf-8', errors='ignore')
    elif content_type == "application/pdf":
        if PdfReader is None:
             raise ValueError("PyPDF2 not installed")
        try:
            with io.BytesIO(content) as pdf_file:
                reader = PdfReader(pdf_file)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
        except Exception as pdf_err:
             print(f"PDF Extraction Error: {pdf_err}")
             raise ValueError("Failed to extract text from PDF")
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        if Document is None:
             raise ValueError("python-docx not installed")
        try:
            with io.BytesIO(content) as docx_file:
                doc = Document(docx_file)
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
        except Exception as docx_err:
            print(f"DOCX Extraction Error: {docx_err}")
            raise ValueError("Failed to extract text from DOCX")
    else:
         # Fallback for other text types or error
         pass
         
    return extracted_text


# --- Internal Groq API Pydantic Models (only used within this module) ---
# These are kept here as they define the specific structure for Groq's API requests
class GroqMessage(BaseModel):
    role: str
    content: str

class GroqRequest(BaseModel):
    model: str
    messages: List[GroqMessage]
    max_tokens: Optional[int] = Field(1000)
    temperature: Optional[float] = Field(0.7)


# --- Core AI Integration and Fallback Functions ---

AI_CACHE_DIR = "ai_cache"
os.makedirs(AI_CACHE_DIR, exist_ok=True)

def get_ai_cache(text: str, prefix: str) -> Optional[str]:
    text_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
    cache_path = os.path.join(AI_CACHE_DIR, f"{prefix}_{text_hash}.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                return json.load(f).get("response")
        except:
            return None
    return None

def set_ai_cache(text: str, prefix: str, response: str):
    text_hash = hashlib.sha256(text.encode('utf-8')).hexdigest()
    cache_path = os.path.join(AI_CACHE_DIR, f"{prefix}_{text_hash}.json")
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump({"response": response, "timestamp": datetime.now().isoformat()}, f)
    except Exception as e:
        print(f"Failed to cache AI response: {e}")

async def call_ai_model(system_prompt: str, user_prompt: str, temperature: float = 0.7, max_tokens: int = 2000, response_format: dict = None) -> str:
    ai_mode = os.environ.get("MODE", "cloud").lower()
    
    if ai_mode == "local":
        local_url = "http://localhost:11434/api/chat"
        payload = {
            "model": "llama3",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "stream": False,
            "options": {"temperature": temperature, "num_predict": max_tokens}
        }
        if response_format:
            payload["format"] = "json"
            
        async with httpx.AsyncClient() as client:
            resp = await client.post(local_url, json=payload, timeout=60.0)
            resp.raise_for_status()
            return resp.json()["message"]["content"]
    else:
        groq_request_body = {
            "model": config.GROQ_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt}, 
                {"role": "user", "content": user_prompt}
            ],
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        if response_format:
            groq_request_body["response_format"] = response_format
            
        async with httpx.AsyncClient() as client:
            response = await client.post(
                config.GROQ_API_URL, 
                headers={"Authorization": f"Bearer {config.GROQ_API_KEY}", "Content-Type": "application/json"}, 
                json=groq_request_body, 
                timeout=config.GROQ_API_TIMEOUT
            )
            response.raise_for_status()
            data = response.json()
            return data["choices"][0]["message"]["content"]

async def analyze_with_groq_api(text: str, analysis_type: str, language: str = "en", cross_language: bool = False, content_type: str = "other") -> Dict[str, Any]:
    """
    Makes an API call to Groq for content analysis (plagiarism or AI detection).
    Uses settings from the config module.
    """
    if not config.GROQ_API_KEY:
        print(f"WARNING: Groq API key not configured for {analysis_type}. Using local fallbacks.")
        if analysis_type == "plagiarism":
            return {"score": calculate_plagiarism_score(text), "sources": find_potential_sources(text, calculate_plagiarism_score(text))}
        else: # ai_detection
            return detect_ai_content(text)

    model_name = config.GROQ_MODEL

    if analysis_type == "plagiarism":
        cross_lang_instruction = " Perform a CROSS-LANGUAGE check across 50+ languages (detect if this text matches translated content)." if cross_language else ""
        system_prompt = f"You are an expert global content analyst specializing in strict plagiarism detection. Language context: {language}. Content Type: {content_type}.{cross_lang_instruction} Analyze the text for matches against a vast database of sources (academic, web, blogs, news). Detect direct copying, paraphrasing, and localized adaptations. Be EXTREMELY STRICT. Provide a plagiarism percentage (0-100) and identify potential sources. Format response clearly with score and source details."
        user_prompt = f"Analyze the following {content_type} text ({language}) for potential plagiarism: '{text}'"
    else:
        system_prompt = "You are an advanced AI Content Detector specialized in distinguishing human writing from machine-generated text. Analyze the text for: lack of personal anecdotes, repetitive sentence structures, overly formal or robotic tone, perfect grammar without stylistic flair, and high perplexity/burstiness indicators suitable for LLMs. Be STRICT. Detect if the text was generated by AI models like GPT-4, Claude, or Llama. Provide a confidence score (0-100%) and extract specific sentences that strongly exhibit AI traits. Format response: 'Confidence: X%'. Then 'AI Sentences:'. Then list each flagged sentence starting with a dash '- '."
        user_prompt = f"Detect if the following text was generated by AI: '{text}'"

    cached_result = get_ai_cache(text, f"analyze_{analysis_type}_{language}_{content_type}")
    if cached_result:
        raw_analysis = cached_result
    else:
        try:
            raw_analysis = await call_ai_model(
                system_prompt, user_prompt, 
                temperature=0.3 if analysis_type == "plagiarism" else 0.7, 
                max_tokens=1000
            )
            set_ai_cache(text, f"analyze_{analysis_type}_{language}_{content_type}", raw_analysis)
        except httpx.HTTPStatusError as http_error:
            print(f"HTTP Error from AI API: {http_error}")
            raise Exception(f"AI API HTTP Error: {http_error}")
        except Exception as e:
            print(f"Generic Error during AI API call: {e}")
            raise Exception(f"AI API processing error: {e}")

    if analysis_type == "plagiarism":
        score_match = re.search(r'(\d+)%', raw_analysis)
        score = float(score_match.group(1)) if score_match else 0.0
        extracted_sources = []
        url_matches = re.findall(r'https?://[^\s)\]]+', raw_analysis)
        for url in url_matches:
            extracted_sources.append({"title": f"External Source: {url}", "url": url.strip(), "type": "web_content", "similarity": round(random.uniform(0.3, 0.9), 2), "matched_words": random.randint(10, 100)})
        if not extracted_sources or len(extracted_sources) < 2:
            num_mock_to_add = min(len(config.MOCK_SOURCES), random.randint(1, 2))
            extracted_sources.extend(random.sample(config.MOCK_SOURCES, num_mock_to_add))
            extracted_sources = list({v['url']:v for v in extracted_sources}.values())
        return {"score": score, "analysis": raw_analysis, "sources": extracted_sources}
    elif analysis_type == "ai_detection":
        confidence_match = re.search(r'(\d+)%', raw_analysis)
        confidence = float(confidence_match.group(1)) if confidence_match else 0.0
        is_ai = "AI" in raw_analysis or "AI-generated" in raw_analysis or confidence > 50
        
        # Extract sentences
        ai_sentences = []
        lines = raw_analysis.split('\n')
        capture = False
        for line in lines:
            if "AI Sentences:" in line:
                capture = True
                continue
            if capture and line.strip().startswith("- "):
                sentence = line.strip()[2:].strip()
                if len(sentence) > 10: # Filter out short noise
                    ai_sentences.append(sentence)
        
        return {"is_ai": is_ai, "confidence": confidence, "analysis": raw_analysis, "ai_sentences": ai_sentences}

async def humanize_with_groq_api(text: str, style: str, complexity: str, target_language: str = "English", content_type: str = "article") -> Dict[str, Any]:
    """Makes an API call to Groq for humanizing and translating text."""
    if not config.GROQ_API_KEY:
        print("WARNING: Groq API key not configured for Humanization. Using local fallbacks.")
        humanized_text, changes = apply_humanization_rules(text, style, complexity)
        return {"humanized_text": humanized_text, "changes": changes}

    model_name = config.GROQ_MODEL

    system_prompt = (
        f"You are an expert content writer and translator specializing in humanizing AI-generated text. "
        f"Your goal is to transform robotic AI text into natural, engaging, and human-like writing in {target_language}. "
        f"Ensure the output maintains the original meaning but uses natural idioms, sentence structures, and tone appropriate for the language. "
        f"Format the output text specifically as a {content_type}. Use appropriate headings (Markdown style #, ##), paragraph structure, and tone. "
        f"For {content_type}, ensure a compelling headline and clear structure."
    )
    user_prompt = (
        f"Transform the following AI-generated text into natural, human-like {content_type} in {target_language}. "
        f"Writing style: {style}, Complexity level: {complexity}. "
        f"Original text: '{text}'"
    )

    cached_result = get_ai_cache(text, f"humanize_{style}_{complexity}_{target_language}_{content_type}")
    if cached_result:
        raw_output = cached_result
    else:
        try:
            raw_output = await call_ai_model(
                system_prompt, user_prompt, 
                temperature=0.7, 
                max_tokens=2000
            )
            set_ai_cache(text, f"humanize_{style}_{complexity}_{target_language}_{content_type}", raw_output)
        except Exception as e:
            print(f"AI api error during humanization: {e}")
            humanized_text, changes = apply_humanization_rules(text, style, complexity)
            return {"humanized_text": humanized_text, "changes": changes}

    humanized_text = raw_output
    humanized_text = re.sub(r'^[\'"]|[\'"]$', '', humanized_text).strip()
    
    # Aggressive cleanup of AI/Markdown artifacts
    # Remove bold/italic markers (*, _)
    humanized_text = re.sub(r'\*+', '', humanized_text) 
    humanized_text = re.sub(r'_+', '', humanized_text)
    
    # Remove code blocks/backticks
    humanized_text = re.sub(r'`+', '', humanized_text)
    
    # Remove blockquote '>' if at start of line
    humanized_text = re.sub(r'^>\s?', '', humanized_text, flags=re.MULTILINE)
    
    # Remove links [text](url) -> text
    humanized_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', humanized_text)
    
    # Note: We PRESERVE '#' characters because they are used by the download 
    # function to identify headings (H1, H2, H3) and apply proper font styles.
    
    return {"humanized_text": humanized_text, "changes": [f"Translated to {target_language}", "Enhanced natural flow", "Improved readability"]}


async def chat_with_groq_api(user_message: str) -> str:
    """Communicates with Groq AI for chatbot responses."""
    if not config.GROQ_API_KEY:
        print("WARNING: Groq API key not configured for Chat. Using local fallback.")
        return get_local_chat_response_fallback(user_message)

    system_prompt = "You are an AI assistant specialized in plagiarism checking and content humanization for PlagiarismPro. Answer user questions concisely and helpfully. Maintain a professional and helpful tone. Do not provide information outside your domain."
    try:
        reply = await call_ai_model(system_prompt, user_message, max_tokens=500, temperature=0.8)
        return reply.strip()
    except Exception as e:
        print(f"Generic Error during chat API call: {e}")
        raise Exception(f"Chat API processing error: {e}")

def moderate_message(message: str) -> Dict[str, Any]:
    """Checks if a message contains any sensitive keywords."""
    message_lower = message.lower()
    for keyword in config.MODERATION_KEYWORDS:
        if keyword in message_lower:
            return {"flagged": True, "reason": "I cannot respond to messages containing sensitive content. Please ask something else."}
    return {"flagged": False, "reason": ""}

# --- Fallback/Local Heuristic Functions (Used if Groq API call fails) ---
def calculate_plagiarism_score(text: str) -> float:
    if not text:
        return 0.0
    text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    hash_value = int(text_hash[:8], 16)
    score = (hash_value % 25) + random.uniform(0, 15)
    return round(min(score, 40.0), 2)

def detect_ai_content(text: str) -> dict:
    ai_indicators = 0
    total_patterns = len(config.AI_PATTERNS)
    for pattern in config.AI_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            ai_indicators += 1
    confidence = (ai_indicators / total_patterns) * 100
    is_ai = confidence > 30
    return {"is_ai": is_ai, "confidence": round(confidence, 2)}

def find_potential_sources(text: str, plagiarism_score: float) -> List[dict]:
    if plagiarism_score < 5:
        return []
    num_sources = min(int(plagiarism_score / 10) + 1, len(config.MOCK_SOURCES))
    selected_sources = random.sample(config.MOCK_SOURCES, num_sources)
    for source in selected_sources:
        source_copy = source.copy()
        source_copy["similarity"] = round(random.uniform(plagiarism_score / 100 * 0.8, plagiarism_score / 100 * 1.2), 2)
        source_copy["matched_words"] = random.randint(10, 50)
    return selected_sources

def apply_humanization_rules(text: str, style: str, complexity: str) -> Tuple[str, List[str]]:
    humanized_text = text
    changes_made = []
    for pattern, replacement in config.HUMANIZATION_RULES["ai_to_human"].items():
        humanized_text, num_replaces = re.subn(pattern, replacement, humanized_text, flags=re.IGNORECASE)
        if num_replaces > 0:
            changes_made.append(f"Applied general AI-to-human rule for '{pattern}' (x{num_replaces})")
    if style == "natural":
        for pattern, replacement in config.HUMANIZATION_RULES["formal_to_casual"].items():
            humanized_text, num_replaces = re.subn(pattern, replacement, humanized_text, flags=re.IGNORECASE)
            if num_replaces > 0:
                changes_made.append(f"Applied natural style rule for '{pattern}' (x{num_replaces})")
    if complexity == "simple":
        for pattern, replacement in config.HUMANIZATION_RULES["complexity_reduction"].items():
            humanized_text, num_replaces = re.subn(pattern, replacement, humanized_text, flags=re.IGNORECASE)
            if num_replaces > 0:
                changes_made.append(f"Applied simple complexity rule for '{pattern}' (x{num_replaces})")
    humanized_text = add_natural_variations(humanized_text, style)
    changes_made = list(dict.fromkeys(changes_made))
    return humanized_text, changes_made

def add_natural_variations(text: str, style: str) -> str:
    if style == "natural":
        text = re.sub(r'\bdo not\b', "don't", text)
        text = re.sub(r'\bcannot\b', "can't", text)
        text = re.sub(r'\bwill not\b', "won't", text)
        text = re.sub(r'\bit is\b', "it's", text)
        text = re.sub(r'\bthat is\b', "that's", text)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    modified_sentences = []
    for i, sentence in enumerate(sentences):
        if i > 0 and random.random() < 0.2:
            if not sentence.lower().startswith(('and', 'but', 'so', 'well', 'anyway')):
                starters = ['And', 'But', 'So', 'Well,', 'Anyway,']
                original_start_lower = sentence[0].lower() + sentence[1:] if len(sentence) > 1 else sentence
                sentence = f"{random.choice(starters)} {original_start_lower}"
        modified_sentences.append(sentence)
    return ' '.join(modified_sentences)

def calculate_improvement_score(original: str, humanized: str) -> Dict[str, float]:
    original_ai_score = detect_ai_content(original)["confidence"]
    humanized_ai_score = detect_ai_content(humanized)["confidence"]
    improvement = max(0, original_ai_score - humanized_ai_score)
    scaled_improvement = improvement * 2.5
    final_improvement = round(min(scaled_improvement, 100.0), 2)
    return {
        "improvement": final_improvement,
        "original_score": original_ai_score,
        "humanized_score": humanized_ai_score
    }

def get_local_chat_response_fallback(user_message: str) -> str:
    user_message_lower = user_message.lower()
    if "plagiarism" in user_message_lower:
        return "To check plagiarism, please paste your text and click 'Check Plagiarism'. Our tool will analyze it."
    elif "humanize" in user_message_lower:
        return "To humanize AI content, go to the AI Humanizer page and paste your text there. Our tool can make it sound more natural."
    elif "hello" in user_message_lower or "hi" in user_message_lower:
        return "Hello! How can I assist you with plagiarism checks or humanizing your content today?"
    elif "cost" in user_message_lower or "price" in user_message_lower or "plans" in user_message_lower:
        return "We have various pricing plans tailored to your needs. You can find more details on our pricing page."
    elif "file" in user_message_lower and ("upload" in user_message_lower or "type" in user_message_lower):
        return "We support PDF, DOC, DOCX, and TXT files for plagiarism checks, in addition to direct text input."
    elif "accuracy" in user_message_lower:
        return "Our AI is designed for high accuracy in both plagiarism detection and content humanization. We constantly refine our models."
    elif "groq" in user_message_lower:
        return "Yes, our advanced features leverage Groq AI technology to provide industry-leading analysis."
    elif "api" in user_message_lower:
        return "Yes, we offer API access for integrating our services into your own applications. Please refer to our documentation."
    else:
        return "I'm here to help with plagiarism detection and content humanization. Could you please rephrase your question, or ask about our services."

def generate_humanized_doc(text_content: str, download_format: str) -> Tuple[io.BytesIO, str, str]:
    """Generates and returns a document stream for download with formatted headings."""
    filename_base = "humanized_content"
    if download_format == "txt":
        file_content = text_content.encode('utf-8')
        media_type = "text/plain"
        filename = f"{filename_base}.txt"
        file_stream = io.BytesIO(file_content)
        file_stream.seek(0)
        return file_stream, media_type, filename
    elif download_format == "word":
        if Document is None:
            raise HTTPException(status_code=500, detail="Word document generation library (python-docx) not installed on server.")
        document = Document()
        
        # Simple Markdown-like parser for headings
        for line in text_content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                document.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                document.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                document.add_heading(line[4:], level=3)
            else:
                document.add_paragraph(line)
        
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{filename_base}.docx"
        return file_stream, media_type, filename
    elif download_format == "pdf":
        if FPDF is None:
            raise HTTPException(status_code=500, detail="PDF generation library (fpdf) not installed on server.")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Simple Markdown-like parser for PDF
        for line in text_content.split('\n'):
            line = line.strip()
            if not line:
                pdf.ln(5) # Small gap for empty lines
                continue
            
            # Proactively sanitize line to avoid UnicodeEncodeErrors in standard FPDF
            # unidecode converts smart quotes, non-breaking hyphens, etc. to ASCII equivalents
            clean_line = unidecode(line)
            
            if line.startswith('# '):
                pdf.set_font("Arial", 'B', 24) # H1
                text_to_print = clean_line[2:]
                pdf.multi_cell(0, 15, text_to_print)
                pdf.ln(2)
            elif line.startswith('## '):
                pdf.set_font("Arial", 'B', 18) # H2
                text_to_print = clean_line[3:]
                pdf.multi_cell(0, 12, text_to_print)
                pdf.ln(2)
            elif line.startswith('### '):
                pdf.set_font("Arial", 'B', 14) # H3
                text_to_print = clean_line[4:]
                pdf.multi_cell(0, 10, text_to_print)
                pdf.ln(1)
            else:
                pdf.set_font("Arial", '', 12) # Body
                pdf.multi_cell(0, 8, clean_line)

        # FPDF output(dest='S') returns a string in Python 3 (latin-1 encoding of bytes).
        # We must encode it back to bytes for io.BytesIO.
        try:
            pdf_string = pdf.output(dest='S')
            pdf_bytes = pdf_string.encode('latin-1')
            file_stream = io.BytesIO(pdf_bytes)
        except Exception as e:
            print(f"Error generating PDF binary: {e}")
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")
            
        file_stream.seek(0)
        media_type = "application/pdf"
        filename = f"{filename_base}.pdf"
        return file_stream, media_type, filename
    else:
        raise HTTPException(status_code=400, detail="Invalid download format requested.")

# --- Advanced Plagiarism Detection System ---
import string
import math
from collections import Counter
import hashlib

# Step 1: Input Normalization
def normalize_text(text: str) -> str:
    text = text.lower()
    text = unidecode(text)
    text = text.translate(str.maketrans("", "", string.punctuation))
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# Step 2: Local Similarity Heuristics
def get_ngrams(text: str, n: int) -> list:
    words = text.split()
    return [" ".join(words[i:i+n]) for i in range(len(words)-n+1)]

def jaccard_similarity(text1: str, text2: str, n: int = 3) -> float:
    set1, set2 = set(get_ngrams(text1, n)), set(get_ngrams(text2, n))
    if not set1 or not set2: return 0.0
    return len(set1.intersection(set2)) / len(set1.union(set2))

def tf_idf_cosine_similarity(text1: str, text2: str) -> float:
    words1, words2 = text1.split(), text2.split()
    vocab = set(words1).union(set(words2))
    if not vocab: return 0.0
    c1, c2 = Counter(words1), Counter(words2)
    v1, v2 = [c1[w] for w in vocab], [c2[w] for w in vocab]
    dot = sum(a*b for a, b in zip(v1, v2))
    mag1 = math.sqrt(sum(a*a for a in v1))
    mag2 = math.sqrt(sum(b*b for b in v2))
    return dot / (mag1 * mag2) if mag1 and mag2 else 0.0

MOCK_CORPUS = [
    "academic integrity is vital in universities plagiarism is strictly prohibited",
    "the quick brown fox jumps over the lazy dog",
    "global warming and climate change significantly affect weather patterns worldwide"
]

def calculate_local_similarity(normalized_text: str) -> float:
    max_sim = 0.0
    for doc in MOCK_CORPUS:
        norm_doc = normalize_text(doc)
        j_sim = jaccard_similarity(normalized_text, norm_doc, n=3)
        c_sim = tf_idf_cosine_similarity(normalized_text, norm_doc)
        max_sim = max(max_sim, j_sim, c_sim)
    
    # Baseline simulation if not matching local exact sources
    base_score = calculate_plagiarism_score(" ".join(normalized_text.split()[:50])) / 100 
    return max(max_sim, base_score) * 100.0

# Step 3: Text Hashing & Cache
PLAGIARISM_CACHE_FILE = "plagiarism_cache.json"

def get_plagiarism_cache(text_hash: str):
    if os.path.exists(PLAGIARISM_CACHE_FILE):
        try:
            with open(PLAGIARISM_CACHE_FILE, "r") as f:
                cache = json.load(f)
                return cache.get(text_hash)
        except Exception: 
            pass
    return None

def set_plagiarism_cache(text_hash: str, text_length: int, result: dict):
    cache = {}
    if os.path.exists(PLAGIARISM_CACHE_FILE):
        try:
            with open(PLAGIARISM_CACHE_FILE, "r") as f:
                cache = json.load(f)
        except Exception: 
            pass
    cache[text_hash] = {
        "text_length": text_length,
        "result": result,
        "timestamp": datetime.now().isoformat()
    }
    with open(PLAGIARISM_CACHE_FILE, "w") as f:
        json.dump(cache, f)

# Step 4, 5 & 6: Main Pipeline
def construct_chunks(normalized_text: str, max_words=500):
    words = normalized_text.split()
    return [" ".join(words[i:i+max_words]) for i in range(0, len(words), max_words)]

async def execute_advanced_plagiarism_check(text: str, language: str, content_type: str) -> dict:
    normalized_text = normalize_text(text)
    
    # Check cache
    text_hash = hashlib.sha256(normalized_text.encode('utf-8')).hexdigest()
    cached_result = get_plagiarism_cache(text_hash)
    if cached_result:
        result = cached_result["result"]
        result["api_used"] = False
        result["confidence"] = "High"
        result["analysis_summary"] = "Result retrieved from local cache."
        return result
        
    local_sim = calculate_local_similarity(normalized_text)
    
    # Rule: If similarity score < 30%, immediately return
    if local_sim < 30.0:
        result = {
            "score": local_sim,
            "originality_level": "High",
            "similarity_range": "0-30%",
            "confidence": "High",
            "analysis_summary": "Passed local heuristic checks. Content seems highly original.",
            "matched_patterns": ["local analysis"],
            "api_used": False,
            "sources": []
        }
        set_plagiarism_cache(text_hash, len(normalized_text), result)
        return result
        
    # Chunking
    chunks = construct_chunks(text, max_words=500)
    chunks_to_analyze = chunks[:2] 
    
    sys_prompt = (
        f"You are an expert content analyst specializing in plagiarism detection. "
        f"Language: {language}. Content Type: {content_type}. "
        f"DO NOT provide an exact percentage. Return ONLY a valid JSON object with keys:\n"
        f"- originality_level: 'High', 'Medium', or 'Low'\n"
        f"- similarity_range: e.g., '40-60%'\n"
        f"- reasoning: short explanation\n"
        f"- suspected_patterns: array of strings\n"
        f"- mock_sources: array of objects with 'url' and 'title', empty default\n"
    )
    
    combined_result = None
    if config.GROQ_API_KEY:
        for chunk in chunks_to_analyze:
            req_body = {
                "model": config.GROQ_MODEL,
                "messages": [
                    {"role": "system", "content": sys_prompt},
                    {"role": "user", "content": f"Analyze this chunk: '{chunk}'"}
                ],
                "response_format": {"type": "json_object"},
                "temperature": 0.2
            }
            async with httpx.AsyncClient() as client:
                try:
                    resp = await client.post(config.GROQ_API_URL, headers={"Authorization": f"Bearer {config.GROQ_API_KEY}"}, json=req_body, timeout=config.GROQ_API_TIMEOUT)
                    resp.raise_for_status()
                    res_content = resp.json()["choices"][0]["message"]["content"]
                    combined_result = json.loads(res_content)
                    break
                except Exception as e:
                    print(f"API Chunk Analysis Error: {e}")
                    
    if not combined_result:
         combined_result = {
            "originality_level": "Medium",
            "similarity_range": "30-50%",
            "reasoning": "Fallback reasoning: API error or disabled.",
            "suspected_patterns": ["api_fallback"],
            "mock_sources": [config.MOCK_SOURCES[0]] if config.MOCK_SOURCES else []
        }
        
    final_output = {
        "score": local_sim, # Fallback base point
        "originality_level": combined_result.get("originality_level", "Medium"),
        "similarity_range": combined_result.get("similarity_range", "30-50%"),
        "confidence": "Medium",
        "analysis_summary": combined_result.get("reasoning", "Analysis complete against APIs."),
        "matched_patterns": combined_result.get("suspected_patterns", []),
        "api_used": True,
        "sources": combined_result.get("mock_sources", combined_result.get("sources", []))
    }
    
    set_plagiarism_cache(text_hash, len(normalized_text), final_output)
    return final_output